#!/usr/bin/env python3
# Python Word Document Text Extractor Service
# Extracts text from Word documents (.doc, .docx) using python-docx and python-docx2txt
import sys
import json
import os

try:
    import docx
    import docx2txt
    from docx import Document
except ImportError as e:
    # When imported as module, don't exit, just raise the error
    if __name__ == "__main__":
        print(json.dumps({"success": False, "error": f"Missing required Python packages: {str(e)}", "text": "", "pages": 0, "metadata": {}, "word_count": 0, "character_count": 0}))
        sys.exit(1)
    else:
        raise ImportError(f"Missing required Python packages: {str(e)}")

def extract_text_with_docx(filepath):
    """Extract text using python-docx library"""
    try:
        doc = Document(filepath)
        text = ""
        paragraphs = []
        
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                paragraphs.append(paragraph.text.strip())
                text += paragraph.text + "\n"
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text += cell.text + " "
                text += "\n"
        
        return {
            "text": text,
            "paragraphs": len(paragraphs),
            "method": "python-docx",
            "success": True
        }
    except Exception as e:
        return {"text": "", "paragraphs": 0, "method": "python-docx", "success": False, "error": str(e)}

def extract_text_with_docx2txt(filepath):
    """Extract text using docx2txt library (simpler approach)"""
    try:
        text = docx2txt.process(filepath)
        paragraphs = len([p for p in text.split('\n') if p.strip()])
        
        return {
            "text": text,
            "paragraphs": paragraphs,
            "method": "docx2txt",
            "success": True
        }
    except Exception as e:
        return {"text": "", "paragraphs": 0, "method": "docx2txt", "success": False, "error": str(e)}

def extract_metadata_from_docx(filepath):
    """Extract metadata from Word document"""
    try:
        doc = Document(filepath)
        metadata = {}
        
        # Get document properties
        core_props = doc.core_properties
        metadata.update({
            "title": core_props.title or "",
            "author": core_props.author or "",
            "subject": core_props.subject or "",
            "keywords": core_props.keywords or "",
            "comments": core_props.comments or "",
            "created": str(core_props.created) if core_props.created else "",
            "modified": str(core_props.modified) if core_props.modified else "",
            "last_modified_by": core_props.last_modified_by or "",
            "revision": core_props.revision or 0
        })
        
        return metadata
    except Exception as e:
        return {"error": str(e)}

def extract_text_from_word(filepath):
    """Extract text from Word document using multiple methods"""
    results = []
    
    # Try different extraction methods
    methods = [
        extract_text_with_docx,
        extract_text_with_docx2txt
    ]
    
    for method in methods:
        result = method(filepath)
        results.append(result)
        
        # If we got good text, use it
        if result["success"] and result["text"].strip() and len(result["text"].strip()) > 10:
            break
    
    # Choose the best result
    successful_results = [r for r in results if r["success"]]
    if not successful_results:
        return {
            "success": False,
            "text": "",
            "paragraphs": 0,
            "metadata": {},
            "word_count": 0,
            "character_count": 0,
            "error": "All extraction methods failed"
        }
    
    best_result = max(successful_results, key=lambda x: len(x["text"].strip()))
    
    # Get metadata
    metadata = extract_metadata_from_docx(filepath)
    
    # Count words and characters
    text = best_result["text"]
    word_count = len(text.split())
    character_count = len(text)
    
    return {
        "success": True,
        "text": text,
        "paragraphs": best_result["paragraphs"],
        "metadata": metadata,
        "word_count": word_count,
        "character_count": character_count,
        "extraction_method": best_result["method"],
        "all_methods": [{"method": r["method"], "success": r["success"], "text_length": len(r["text"]), "error": r.get("error")} for r in results]
    }

if __name__ == "__main__":
    if len(sys.argv) < 2:
        print(json.dumps({"success": False, "error": "No file path provided", "text": "", "paragraphs": 0, "metadata": {}, "word_count": 0, "character_count": 0}))
        sys.exit(1)
    
    word_filepath = sys.argv[1]
    if not os.path.exists(word_filepath):
        print(json.dumps({"success": False, "error": f"File not found: {word_filepath}", "text": "", "paragraphs": 0, "metadata": {}, "word_count": 0, "character_count": 0}))
        sys.exit(1)
    
    result = extract_text_from_word(word_filepath)
    print(json.dumps(result, indent=2))
